import logging
from pathlib import Path
from typing import List

from sqlalchemy.orm import Session

from ..models import File, FulltextBlock

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover - optional import, validated at runtime
    fitz = None  # type: ignore[assignment]

try:
    from docx import Document as DocxDocument  # python-docx
except Exception:  # pragma: no cover - optional import, validated at runtime
    DocxDocument = None  # type: ignore[assignment]


logger = logging.getLogger(__name__)


class FulltextService:
    """负责从 File.storage_path 提取全文文本并缓存到 FulltextBlock。

    - PDF: 使用 PyMuPDF 按页抽取文本。
    - Word: 使用 python-docx 抽取段落与表格。
    - Image: 目前占位实现，后续可接入整页 OCR。
    """

    def __init__(self, db: Session) -> None:
        self.db = db
        self._backend_root = Path(__file__).resolve().parents[2]

    def get_or_extract_fulltext(self, file_id: int) -> List[FulltextBlock]:
        """获取或抽取指定文件的全文文本块列表。

        优先读取缓存的 FulltextBlock；若不存在则根据 source_type 执行一次抽取并落库。
        """

        cached = (
            self.db.query(FulltextBlock)
            .filter(FulltextBlock.file_id == file_id)
            .order_by(FulltextBlock.page_num.asc(), FulltextBlock.block_index.asc())
            .all()
        )
        if cached:
            logger.info("fulltext.get_or_extract: hit cache file_id=%s blocks=%s", file_id, len(cached))
            return cached

        file: File | None = self.db.query(File).filter(File.id == file_id).first()
        if not file:
            raise ValueError(f"File not found: id={file_id}")

        storage_path = (self._backend_root / file.storage_path).resolve()
        if not storage_path.exists():
            raise FileNotFoundError(f"Storage file not found: {storage_path}")

        logger.info(
            "fulltext.get_or_extract: start file_id=%s type=%s path=%s",
            file_id,
            file.source_type,
            storage_path,
        )

        blocks: List[FulltextBlock] = []
        if file.source_type == "pdf":
            for idx, item in enumerate(self._extract_text_from_pdf(storage_path)):
                if not item["content"]:
                    continue
                block = FulltextBlock(
                    file_id=file_id,
                    page_num=item["page_num"],
                    block_index=idx,
                    content=item["content"],
                )
                self.db.add(block)
                blocks.append(block)
        elif file.source_type == "word":
            for idx, item in enumerate(self._extract_text_from_word(storage_path)):
                if not item["content"]:
                    continue
                block = FulltextBlock(
                    file_id=file_id,
                    page_num=item["page_num"],
                    block_index=idx,
                    content=item["content"],
                )
                self.db.add(block)
                blocks.append(block)
        elif file.source_type == "image":
            for idx, item in enumerate(self._extract_text_from_image(storage_path)):
                if not item["content"]:
                    continue
                block = FulltextBlock(
                    file_id=file_id,
                    page_num=item["page_num"],
                    block_index=idx,
                    content=item["content"],
                )
                self.db.add(block)
                blocks.append(block)
        else:
            logger.warning("fulltext.get_or_extract: unsupported source_type=%s", file.source_type)

        if blocks:
            self.db.commit()
            logger.info("fulltext.get_or_extract: stored blocks file_id=%s count=%s", file_id, len(blocks))
        else:
            logger.warning("fulltext.get_or_extract: no text extracted file_id=%s", file_id)

        return blocks

    def _extract_text_from_pdf(self, path: Path) -> List[dict]:
        """使用 PyMuPDF 从 PDF 中按页抽取文本。"""

        if fitz is None:
            raise RuntimeError("PyMuPDF (fitz) is not installed")

        doc = fitz.open(path)  # type: ignore[arg-type]
        items: List[dict] = []
        try:
            for page_index in range(doc.page_count):
                page = doc.load_page(page_index)
                text = page.get_text("text") or ""
                content = text.strip()
                items.append({"page_num": page_index + 1, "content": content})
        finally:
            doc.close()
        return items

    def _extract_text_from_word(self, path: Path) -> List[dict]:
        """使用 python-docx 从 Word 文档中抽取段落与表格内容。"""

        if DocxDocument is None:
            raise RuntimeError("python-docx is not installed")

        doc = DocxDocument(str(path))
        items: List[dict] = []

        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                items.append({"page_num": 1, "content": text})

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    items.append({"page_num": 1, "content": " | ".join(cells)})

        return items

    def _extract_text_from_image(self, path: Path) -> List[dict]:
        """针对图片文件的全文抽取占位实现。

        目前仅返回空列表，后续可以接入整页 OCR（例如复用 OcrService）。
        """

        logger.warning("fulltext._extract_text_from_image: not implemented, path=%s", path)
        return []
